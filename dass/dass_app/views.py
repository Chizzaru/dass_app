from re import template
from django.shortcuts import render
from django.http import HttpResponse, HttpResponseRedirect
from django.template import loader
from .models import Users, Question, Choice, Answer, Result
from .operation import getResultAnxiety, getResultDepression, getResultStress
import datetime


from docx import Document
from docx.shared import Inches

from .forms import DassForm

def index(request):
    myUsers = Users.objects.all().values()
    template = loader.get_template('index.html')
    context = {
        'myUsers' : myUsers,
    }
    return HttpResponse(template.render(context,request))

def form(request):
    myQuestions = Question.objects.all().values()
    myChoices = Choice.objects.all().values()
    template = loader.get_template('form.html')
    context = {
        'myQuestions'   : myQuestions,
        'myChoices'     : myChoices
    }
    return HttpResponse(template.render(context, request))

#def result(request):
#    template = loader.get_template('error.html')
#    return HttpResponse(template.render())

def submitform(request):
    # if this is a POST request we need to process the form data
    if request.method == 'POST':
        counter = 0
        chosen_number = 0
        stname = request.POST['name']
        tdate = datetime.datetime.now()

        #mydata = Answer.objects.filter(lastname=stname, date_answered=tdate).values()
        #if not mydata:
        #    return HttpResponseRedirect('error')

        for x in request.POST.getlist('question[]'):
            item_question = Question.objects.get(id=x)
            choice = request.POST[x]
            item_choose = Choice.objects.get(id=choice)
            
            for_save = Answer(name=stname,question=item_question,answer=item_choose,date_answered=tdate)
            for_save.save()
            counter += 1
            chosen_number += int(choice)
        
        dscore = chosen_number / counter
        score_list = [getResultDepression(dscore), getResultAnxiety(dscore), getResultStress(dscore)]
        context = {
            'score' : score_list
        }

        save_score = Result(name=stname, score = dscore, depression=getResultDepression(dscore),anxiety=getResultAnxiety(dscore),stress=getResultStress(dscore),date=tdate)
        save_score.save()

        #return HttpResponseRedirect('form')

        #create docx file

        """ document = Document()

        document.add_heading('Depression Anxiety Stress Scale Result', 0)

        p = document.add_paragraph('Your DASS Score : ' + str(dscore))
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        #p.add_run('italic.').italic = True

        a = getResultDepression(dscore)
        b = getResultAnxiety(dscore)
        c = getResultStress(dscore)

        

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Depression'
        hdr_cells[1].text = 'Anxiety'
        hdr_cells[2].text = 'Stress'
        row_cells = table.add_row().cells
        row_cells[0].text = a
        row_cells[1].text = b
        row_cells[2].text = c

        document.add_page_break()

        document.save('dass.docx') """


        #end creating docx

        template = loader.get_template('result.html')
        return HttpResponse(template.render(context, request))

    

def get_name(request):
    # if this is a POST request we need to process the form data
    if request.method == 'POST':
        # create a form instance and populate it with data from the request:
        form = DassForm(request.POST)
        # check whether it's valid:
        if form.is_valid():
            # process the data in form.cleaned_data as required
            # ...
            # redirect to a new URL:
            return HttpResponseRedirect('/thanks/')

    # if a GET (or any other method) we'll create a blank form
    else:
        form = DassForm()

    return render(request, 'dassform.html', {'form': form})